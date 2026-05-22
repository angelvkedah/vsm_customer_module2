import io
import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

from modules.vsm_protocol.handlers.human_readable import get_human_message_templates
from utils.helpers import clean_text, format_car_number, format_datetime, safe_str


STATIC_DIR = Path(__file__).resolve().parent.parent / "static"
LOGO_PATH = STATIC_DIR / "vsm_service_logo.png"

COMPANY_NAME = "ООО «ВСМ-Сервис»"
DOC_FONT_NAME = "Arial"


def get_protocol_message_text(row):
    code = safe_str(row.get("messagecode", ""))
    event_type = safe_str(row.get("event_type", ""))
    message_text = safe_str(row.get("message_text", ""))

    templates = get_human_message_templates(code)

    if event_type == "activation":
        text = templates.get("kurztext_3") or message_text
    elif event_type == "deactivation":
        text = templates.get("kurztext_4") or message_text
    else:
        text = message_text or templates.get("kurztext_2", "")

    text = clean_text(text)

    text = re.sub(
        rf"\bДС\s+{re.escape(code)}\b",
        f"ДС [{code}]",
        text
    )

    if f"[{code}]" in text:
        return text

    return f"[{code}] {text}"


def build_human_readable_entry(row):
    train_id = safe_str(row.get("train_id", ""))
    carnumber = format_car_number(row.get("carnumber", ""))
    timestamp = row.get("timestamp", None)

    timestamp_str = format_datetime(timestamp).split(" ")[1] if timestamp else ""
    message_text = get_protocol_message_text(row)

    lines = []

    header_parts = []

    if train_id:
        header_parts.append(f"{train_id}")

    if carnumber:
        header_parts.append(f"{carnumber}")

    if header_parts:
        lines.append(", ".join(header_parts) + ".")

    if timestamp_str and message_text:
        lines.append(f"{timestamp_str} {message_text}.")
    elif timestamp_str:
        lines.append(f"{timestamp_str} Зафиксировано диагностическое сообщение.")

    return "\n".join(lines)


def build_human_readable_protocol_text(timeline_df, train_human_name, dt_from, dt_to):
    lines = [
        "Эксплуатационный протокол",
        "",
        f"Поезд: {train_human_name}",
        f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}",
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
    ]

    if timeline_df is None or timeline_df.empty:
        lines.append("За указанный период диагностические события не обнаружены.")
        return "\n".join(lines)

    for _, row in timeline_df.iterrows():
        lines.append(build_human_readable_entry(row))
        lines.append("")

    return "\n".join(lines).strip()


def export_text_to_docx(protocol_text, file_title="Эксплуатационный протокол"):
    try:
        doc = Document()
        lines = str(protocol_text).splitlines()

        if lines:
            first_line = lines[0].strip()
            if first_line:
                title = doc.add_heading(first_line, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lines = lines[1:]

        for line in lines:
            doc.add_paragraph(clean_text(line)) if line.strip() else doc.add_paragraph("")

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        print(f"EDITED DOCX export error: {e}")
        return None


def get_column_names_map():
    return {
        "train_id": "Номер поезда",
        "carnumber": "Вагон",
        "messagecode": "Код ДС",
        "event_type": "Тип события",
        "timestamp": "Время события",
        "message_text": "Описание",
        "duration_str": "Продолжительность",
        "parsingtime": "Время парсинга",
    }


def prepare_row_for_export(row, col):
    value = row.get(col, "")

    if col == "event_type":
        if value == "activation":
            return "Активация"
        if value == "deactivation":
            return "Деактивация"
        if value == "still_active_marker":
            return "Активно до сих пор"
        return "—"

    if col == "timestamp":
        return format_datetime(value)

    return clean_text(value)


def export_to_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        doc = Document()

        setup_document_styles(doc)
        add_protocol_header(doc)
        add_protocol_footer(doc)

        add_subject_block(
            doc,
            train_human_name,
            dt_from,
            dt_to
        )

        add_protocol_table(
            doc,
            timeline_df,
            selected_columns=selected_columns
        )

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        print(f"DOCX export error: {e}")
        return None


def set_cell_text(cell, text, font_size=7, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = DOC_FONT_NAME
    run.font.size = Pt(font_size)
    run.bold = bold
    return paragraph


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = DOC_FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold

    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_page_number(paragraph):
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_total_pages(paragraph):
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "NUMPAGES"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def setup_document_styles(doc):
    section = doc.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = DOC_FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_NAME)
    style.font.size = Pt(11)


def add_protocol_header(doc):
    header = doc.sections[0].header
    header.is_linked_to_previous = False

    table = header.add_table(rows=1, cols=2, width=Inches(5.2))
    table.autofit = True

    logo_cell = table.cell(0, 0)
    title_cell = table.cell(0, 1)

    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if LOGO_PATH.exists():
        run = logo_paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.1))

    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.paragraph_format.space_before = Pt(14)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title_run = title_paragraph.add_run("Эксплуатационный протокол")
    set_run_font(title_run, size=12, bold=True, color=(0, 0, 0))


def add_protocol_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    paragraph = footer.paragraphs[0]
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    border = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')

    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '7')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '000000')

    border.append(top)
    pPr.append(border)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_company = paragraph.add_run(f"{COMPANY_NAME}.")
    set_run_font(run_company, size=7)

    paragraph.add_run("\t\t")

    run_page = paragraph.add_run("Страница ")
    set_run_font(run_page, size=7)

    add_page_number(paragraph)

    run_of = paragraph.add_run(" из ")
    set_run_font(run_of, size=7)

    add_total_pages(paragraph)

    for run in paragraph.runs:
        run.font.name = DOC_FONT_NAME
        run.font.size = Pt(7)


def add_subject_block(doc, train_human_name, dt_from, dt_to):
    subject = doc.add_paragraph()
    set_paragraph_spacing(subject, before=18, after=14)

    date_part = format_datetime(dt_from).split(" ")[0] if dt_from else ""
    time_from = format_datetime(dt_from).split(" ")[1] if dt_from else ""
    time_to = format_datetime(dt_to).split(" ")[1] if dt_to else ""

    run_value = subject.add_run(
        f"{train_human_name} за период {date_part} {time_from} – {time_to}"
    )

    set_run_font(run_value, size=14, bold=True)


def add_event_paragraph(doc, row):
    entry_text = build_human_readable_entry(row)

    for line in entry_text.split("\n"):
        if not line.strip():
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=1.0)

        run = paragraph.add_run(clean_text(line))
        set_run_font(run, size=12)


def add_no_events_paragraph(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, before=0, after=6)

    run = paragraph.add_run(
        "За указанный период диагностические события не обнаружены."
    )
    set_run_font(run, size=12)

def set_table_cell_shading(cell, fill_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_color)
    tc_pr.append(shading)


def set_table_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def format_table_cell(cell, text, font_size=9, bold=False, align="left"):
    cell.text = ""

    paragraph = cell.paragraphs[0]

    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(clean_text(text))
    set_run_font(run, size=font_size, bold=bold)


def prepare_table_value(row, col):
    if col == "carnumber":
        return format_car_number(row.get(col, ""))

    if col == "messagecode":
        code = safe_str(row.get(col, ""))
        return f"[{code}]" if code else ""

    if col == "message_text":
        return get_protocol_message_text(row)

    return str(prepare_row_for_export(row, col))


def add_protocol_table(doc, timeline_df, selected_columns=None):
    if timeline_df is None or timeline_df.empty:
        add_no_events_paragraph(doc)
        return

    if not selected_columns:
        selected_columns = [
            "train_id",
            "carnumber",
            "messagecode",
            "event_type",
            "timestamp",
            "message_text",
        ]

    column_names = get_column_names_map()
    selected_cols = [
        col for col in selected_columns
        if col in timeline_df.columns
    ]

    if not selected_cols:
        selected_cols = [
            "train_id",
            "carnumber",
            "messagecode",
            "event_type",
            "timestamp",
            "message_text",
        ]
        selected_cols = [
            col for col in selected_cols
            if col in timeline_df.columns
        ]

    headers = [
        column_names.get(col, col)
        for col in selected_cols
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    # Примерные ширины колонок под A4
    width_map = {
        "train_id": 0.65,
        "carnumber": 0.45,
        "messagecode": 0.55,
        "event_type": 0.65,
        "timestamp": 1.15,
        "message_text": 4.35,
        "duration_str": 0.75,
        "parsingtime": 1.10,
    }

    header_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        col = selected_cols[i]
        cell = header_cells[i]

        set_table_cell_shading(cell, "D9D9D9")
        set_table_cell_width(cell, width_map.get(col, 1.0))
        format_table_cell(
            cell,
            header,
            font_size=9,
            bold=True,
            align="center"
        )

    # Повтор заголовка таблицы на новых страницах
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

    for _, row in timeline_df.iterrows():
        cells = table.add_row().cells

        for i, col in enumerate(selected_cols):
            value = prepare_table_value(row, col)
            cell = cells[i]

            set_table_cell_width(cell, width_map.get(col, 1.0))

            align = "left"
            if col in ["messagecode", "event_type", "carnumber"]:
                align = "center"

            format_table_cell(
                cell,
                value,
                font_size=8,
                bold=False,
                align=align
            )

def export_human_readable_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        doc = Document()

        setup_document_styles(doc)
        add_protocol_header(doc)
        add_protocol_footer(doc)

        add_subject_block(
            doc,
            train_human_name,
            dt_from,
            dt_to
        )

        if timeline_df is None or timeline_df.empty:
            add_no_events_paragraph(doc)
        else:
            for _, row in timeline_df.iterrows():
                add_event_paragraph(doc, row)

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        print(f"HUMAN DOCX export error: {e}")
        return None


def export_to_xlsx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Эксплуатационный протокол"

        ws["A1"] = f"Поезд: {train_human_name}"
        ws["A2"] = f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}"
        ws["A3"] = f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"

        if not selected_columns:
            selected_columns = [
                "train_id",
                "carnumber",
                "messagecode",
                "event_type",
                "timestamp",
                "message_text",
            ]

        column_names = get_column_names_map()
        selected_cols = [col for col in selected_columns if col in timeline_df.columns]
        headers = [column_names.get(col, col) for col in selected_cols]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        for row_idx, (_, row) in enumerate(timeline_df.iterrows(), 6):
            for col_idx, col in enumerate(selected_cols, 1):
                ws.cell(row=row_idx, column=col_idx, value=prepare_row_for_export(row, col))

        for col in range(1, len(headers) + 1):
            column_letter = ws.cell(row=5, column=col).column_letter
            ws.column_dimensions[column_letter].width = 25

        xlsx_bytes = io.BytesIO()
        wb.save(xlsx_bytes)
        xlsx_bytes.seek(0)
        return xlsx_bytes

    except Exception as e:
        print(f"XLSX export error: {e}")
        return None


def export_to_csv(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        if not selected_columns:
            selected_columns = [
                "train_id",
                "carnumber",
                "messagecode",
                "event_type",
                "timestamp",
                "message_text",
            ]

        available_columns = [col for col in selected_columns if col in timeline_df.columns]
        df_clean = timeline_df[available_columns].copy()

        if "event_type" in df_clean.columns:
            event_type_map = {
                "activation": "Активация",
                "deactivation": "Деактивация",
                "still_active_marker": "Активно до сих пор",
            }
            df_clean["event_type"] = df_clean["event_type"].map(event_type_map).fillna(df_clean["event_type"])

        if "timestamp" in df_clean.columns:
            df_clean["timestamp"] = df_clean["timestamp"].apply(format_datetime)

        for col in df_clean.columns:
            if df_clean[col].dtype == "object":
                df_clean[col] = df_clean[col].apply(clean_text)

        csv_data = io.StringIO()
        df_clean.to_csv(csv_data, index=False, encoding="utf-8-sig")
        return io.BytesIO(csv_data.getvalue().encode("utf-8-sig"))

    except Exception as e:
        print(f"CSV export error: {e}")
        return None